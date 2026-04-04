import os
import re
import hashlib
import importlib
import json
import subprocess
import sys
import time
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime
from statistics import mean, median
from typing import Any, Dict, List, Mapping, Optional, Tuple, TypedDict, cast
from collections import Counter, defaultdict
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.figure import Figure
import pandas as pd
from dotenv import load_dotenv

try:
    from google import genai  # pyright: ignore[reportMissingImports]
except ImportError:
    genai = cast(Any, None)

try:
    from google.genai import types as genai_types  # pyright: ignore[reportMissingImports]
except ImportError:
    genai_types = cast(Any, None)

try:
    from openai import OpenAI  # pyright: ignore[reportMissingImports]
except ImportError:
    OpenAI = cast(Any, None)


WORKFLOW_STAGES = {
    "batched": "elements_batched",
    "primary_completed": "primary_translation_completed",
    "primary_retry_completed": "primary_retry_completed",
    "evaluation_completed": "evaluation_completed",
    "fallback_completed": "fallback_completed",
    "finalized": "finalized",
}


def build_workflow_metadata(
    stage: Optional[str] = None,
    docxfilename: Optional[str] = None,
    source_language: Optional[str] = None,
    target_language: Optional[str] = None,
    primary_model_name: Optional[str] = None,
    evaluation_model_name: Optional[str] = None,
    fallback_model_name: Optional[str] = None,
    primary_system_message: Optional[str] = None,
    evaluation_system_message: Optional[str] = None,
    fallback_system_message: Optional[str] = None,
) -> Dict[str, Optional[str]]:
    """Build the standard metadata dictionary for workflow checkpoints."""
    return {
        "stage": stage,
        "docxfilename": docxfilename,
        "source_language": source_language,
        "target_language": target_language,
        "primary_model_name": primary_model_name,
        "evaluation_model_name": evaluation_model_name,
        "fallback_model_name": fallback_model_name,
        "primary_system_message": primary_system_message,
        "evaluation_system_message": evaluation_system_message,
        "fallback_system_message": fallback_system_message,
    }


def get_word_files_dir(base_dir: Optional[str] = None) -> str:
    """
    Return the path to the 'word_files' subdirectory rooted at base_dir or CWD.
    """
    root = base_dir or os.getcwd()
    word_dir = os.path.join(root, "word_files")
    return word_dir


def list_docx_files(word_dir: Optional[str] = None) -> List[str]:
    """
    List *.docx files inside the provided word_dir (defaults to 'word_files').
    """
    target_dir = word_dir or get_word_files_dir()
    if not os.path.isdir(target_dir):
        print(f"'word_files' directory not found at {target_dir}")
        return []

    docx_files = sorted(
        fname for fname in os.listdir(target_dir) if fname.lower().endswith(".docx")
    )

    if docx_files:
        print("Found DOCX files:")
        for fname in docx_files:
            print(f"- {fname}")
    else:
        print("No DOCX files found in 'word_files'.")

    return docx_files


import os
from typing import Optional

def verify_docx_file(docx_filename: str, word_dir: Optional[str] = None) -> str:
    """
    Ensure the requested DOCX file exists inside the word_files directory.
    """
    target_dir = word_dir or get_word_files_dir()
    docx_path = os.path.join(target_dir, docx_filename)

    if not os.path.isfile(docx_path):
        raise FileNotFoundError(
            f"{docx_filename} was not found in the 'word_files' directory ({target_dir})."
        )

    base_dir = os.path.dirname(os.path.dirname(target_dir))
    relative_docx_path = os.path.relpath(docx_path, start=base_dir)

    print(f"Confirmed DOCX present: {relative_docx_path}")
    return docx_path

def summarize_docx_styles(docx_path: str):
    """
    Inspect a DOCX file and summarize:
      - Paragraph styles in use
      - Inline run-level styles: bold, italic, underline, superscript, subscript
      - Examples of styled text

    Returns a dictionary with summary info.
    Prints a human-readable overview.
    """

    doc = Document(docx_path)
    display_docx_path = os.path.relpath(docx_path).replace("/", "\\")

    # Counters
    paragraph_style_counts = Counter()
    inline_counts = Counter()
    inline_examples = defaultdict(list)

    for paragraph in doc.paragraphs:
        # Paragraph style
        style_name = paragraph.style.name if paragraph.style else "Unknown"
        paragraph_style_counts[style_name] += 1

        # Inspect runs inside paragraph
        for run in paragraph.runs:
            text = (run.text or "").strip()
            if not text:
                continue

            # Bold
            if run.bold:
                inline_counts["bold"] += 1
                if len(inline_examples["bold"]) < 10:
                    inline_examples["bold"].append(text)

            # Italic
            if run.italic:
                inline_counts["italic"] += 1
                if len(inline_examples["italic"]) < 10:
                    inline_examples["italic"].append(text)

            # Underline
            if run.underline:
                inline_counts["underline"] += 1
                if len(inline_examples["underline"]) < 10:
                    inline_examples["underline"].append(text)

            # Superscript / subscript
            va = getattr(run.font, "vertical_alignment", None)
            if va:
                if "SUPERSCRIPT" in str(va):
                    inline_counts["superscript"] += 1
                    if len(inline_examples["superscript"]) < 10:
                        inline_examples["superscript"].append(text)
                elif "SUBSCRIPT" in str(va):
                    inline_counts["subscript"] += 1
                    if len(inline_examples["subscript"]) < 10:
                        inline_examples["subscript"].append(text)

    # ----- PRINT REPORT -----
    print("──────────────────────────────────────────")
    print(f"Document: {display_docx_path}")
    print("──────────────────────────────────────────")

    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Unique paragraph styles: {len(paragraph_style_counts)}")

    for style, count in paragraph_style_counts.most_common():
        print(f"  {count:4}  {style}")

    print("\nInline formatting counts:")
    if not inline_counts:
        print("  (No inline formatting detected)")
    else:
        for k, v in inline_counts.items():
            print(f"  {k:12}: {v}")

        print("\nInline formatting examples (up to 10 each):")
        for k, examples in inline_examples.items():
            print(f"  {k}:")
            for ex in examples:
                print(f"     - {ex}")

    print("──────────────────────────────────────────")

    # Return structured summary for programmatic use
    return {
        "paragraph_style_counts": paragraph_style_counts,
        "inline_counts": inline_counts,
        "inline_examples": inline_examples,
    }


class WorkflowElement(TypedDict):
    element_number: int
    element_type: str
    word_style: str
    text: str
    element_id: str


class WorkflowRuntimeElement(WorkflowElement, total=False):
    tokens: int
    batch_number: int


def extract_docx_elements(docx_path: str) -> Tuple[List[WorkflowRuntimeElement], str]:
    """
    Extract non-blank DOCX paragraphs and footnotes into JSON-ready elements.
    Returns (elements, summary_string).
    """

    doc = Document(docx_path)
    word_namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def normalize_text(text: str) -> str:
        return text.replace("\xa0", " ")

    def markdown_escape(text: str) -> str:
        # Keep prose readable by escaping only characters that commonly break
        # inline Markdown emphasis/code spans in plain text extraction.
        return re.sub(r"([\\`*_])", r"\\\1", text)

    def run_to_md(run) -> str:
        text = normalize_text(run.text or "")
        if not text:
            return ""

        formatted = markdown_escape(text)
        vertical_alignment = getattr(run.font, "vertical_alignment", None)
        vertical_alignment_text = str(vertical_alignment) if vertical_alignment else ""

        if "SUPERSCRIPT" in vertical_alignment_text:
            formatted = f"<sup>{formatted}</sup>"
        elif "SUBSCRIPT" in vertical_alignment_text:
            formatted = f"<sub>{formatted}</sub>"

        bold = run.bold is True
        italic = run.italic is True
        if bold and italic:
            formatted = f"***{formatted}***"
        elif bold:
            formatted = f"**{formatted}**"
        elif italic:
            formatted = f"*{formatted}*"

        if run.underline is True:
            formatted = f"<u>{formatted}</u>"

        return formatted

    def list_info(paragraph):
        props = paragraph._p.pPr
        if props is None or props.numPr is None:
            return False, None, 0

        level_element = props.numPr.ilvl
        level = int(level_element.val) if level_element is not None else 0
        return True, "1. ", level

    def paragraph_text(paragraph) -> str:
        inline = "".join(run_to_md(run) for run in paragraph.runs)
        if not inline.strip():
            return ""

        is_list, marker, level = list_info(paragraph)
        if is_list:
            indent = " " * (4 * level)
            return f"{indent}{marker}{inline}".rstrip()
        return inline.rstrip()

    def footnote_run_to_md(run_element) -> str:
        text_parts: List[str] = []
        for child in run_element:
            if child.tag == f"{{{word_namespace['w']}}}t":
                text_parts.append(normalize_text(child.text or ""))
            elif child.tag == f"{{{word_namespace['w']}}}tab":
                text_parts.append("\t")
            elif child.tag in {
                f"{{{word_namespace['w']}}}br",
                f"{{{word_namespace['w']}}}cr",
            }:
                text_parts.append("\n")

        text = normalize_text("".join(text_parts))
        if not text:
            return ""

        formatted = markdown_escape(text)
        properties = run_element.find("w:rPr", word_namespace)
        if properties is not None:
            vertical_alignment = properties.find("w:vertAlign", word_namespace)
            vertical_value = (
                vertical_alignment.get(f"{{{word_namespace['w']}}}val", "")
                if vertical_alignment is not None
                else ""
            )
            if vertical_value == "superscript":
                formatted = f"<sup>{formatted}</sup>"
            elif vertical_value == "subscript":
                formatted = f"<sub>{formatted}</sub>"

            bold = properties.find("w:b", word_namespace) is not None
            italic = properties.find("w:i", word_namespace) is not None
            if bold and italic:
                formatted = f"***{formatted}***"
            elif bold:
                formatted = f"**{formatted}**"
            elif italic:
                formatted = f"*{formatted}*"

            if properties.find("w:u", word_namespace) is not None:
                formatted = f"<u>{formatted}</u>"

        return formatted

    def footnote_paragraph_text(paragraph_element) -> str:
        parts: List[str] = []
        for child in paragraph_element:
            if child.tag == f"{{{word_namespace['w']}}}r":
                parts.append(footnote_run_to_md(child))
            elif child.tag == f"{{{word_namespace['w']}}}hyperlink":
                for run_element in child.findall("w:r", word_namespace):
                    parts.append(footnote_run_to_md(run_element))
        return "".join(parts).strip()

    def element_hash(element_type: str, style_name: str, text: str) -> str:
        raw = f"{element_type}|{style_name}|{text}".encode("utf-8")
        return hashlib.sha1(raw).hexdigest()[:12]

    elements: List[WorkflowRuntimeElement] = []
    blank_body_paragraphs = 0
    for paragraph in doc.paragraphs:
        style_name = getattr(paragraph.style, "name", None) or "Unknown"
        text = paragraph_text(paragraph).strip()
        if not text:
            blank_body_paragraphs += 1
            continue

        element_number = len(elements) + 1
        elements.append(
            {
                "element_number": element_number,
                "element_type": "paragraph",
                "word_style": style_name,
                "text": text,
                "element_id": element_hash("paragraph", style_name, text),
            }
        )

    body_count = len(elements)
    footnote_count = 0

    try:
        with zipfile.ZipFile(docx_path) as docx_zip:
            if "word/footnotes.xml" in docx_zip.namelist():
                footnotes_root = ET.fromstring(docx_zip.read("word/footnotes.xml"))
                for footnote in footnotes_root.findall("w:footnote", word_namespace):
                    footnote_type = footnote.get(f"{{{word_namespace['w']}}}type")
                    if footnote_type in {"separator", "continuationSeparator"}:
                        continue

                    for paragraph_element in footnote.findall("w:p", word_namespace):
                        text = footnote_paragraph_text(paragraph_element)
                        if not text:
                            continue

                        style_element = paragraph_element.find("w:pPr/w:pStyle", word_namespace)
                        style_name = (
                            style_element.get(f"{{{word_namespace['w']}}}val")
                            if style_element is not None
                            else "Footnote"
                        ) or "Footnote"
                        element_number = len(elements) + 1
                        elements.append(
                            {
                                "element_number": element_number,
                                "element_type": "footnote",
                                "word_style": style_name,
                                "text": text,
                                "element_id": element_hash("footnote", style_name, text),
                            }
                        )
                        footnote_count += 1
    except (KeyError, zipfile.BadZipFile, ET.ParseError):
        footnote_count = 0

    sample_lines = [
        f"Extracted {body_count} body elements from DOCX.",
        (
            f"Appended {footnote_count} footnote elements."
            if footnote_count
            else "No footnotes found."
        ),
        f"Skipped {blank_body_paragraphs} blank body paragraphs.",
    ]
    if footnote_count:
        sample_lines.append("Footnote preview:")
        for element in elements[body_count:body_count + min(footnote_count, 8)]:
            preview_text = re.sub(r"\\([\\`*_{}\[\]()#+\-.!|>~])", r"\1", element["text"])
            label = f"[{element['element_type']}]"
            if element["word_style"].lower() != element["element_type"].lower():
                label = f"{label} [{element['word_style']}]"
            sample_lines.append(
                f"{element['element_number']}. {label} {preview_text[:120]}"
            )

    summary = "\n".join(sample_lines)
    print(summary)
    return elements, summary


def _ensure_tiktoken():
    """Import tiktoken, installing it if needed."""
    try:
        return importlib.import_module("tiktoken")
    except ModuleNotFoundError:
        print("tiktoken not found; installing...", flush=True)
        subprocess.check_call([sys.executable, "-m", "pip", "install", "tiktoken"])
        return importlib.import_module("tiktoken")


def count_tokens_with_tiktoken(
    elements: List[WorkflowRuntimeElement], model_name: str = "gpt-4o"
) -> str:
    """
    Add token counts for each element and return a status summary string.
    """
    if not isinstance(elements, list):
        raise TypeError("`elements` must be a list of dictionaries.")

    tiktoken_module = _ensure_tiktoken()
    try:
        encoder = tiktoken_module.encoding_for_model(model_name)
    except Exception:
        encoder = tiktoken_module.get_encoding("cl100k_base")

    updated = 0
    for element in elements:
        text = element.get("text", "") or ""
        element["tokens"] = len(encoder.encode(text)) if text else 0
        updated += 1

    token_values = [int(element.get("tokens", 0) or 0) for element in elements]
    if token_values:
        summary = (
            f"Updated {updated} elements with tiktoken counts "
            f"(model: {model_name}) | " 
            f"min={min(token_values)} | median={median(token_values):.0f} | "
            f"max={max(token_values)} | mean={mean(token_values):.1f}"
        )
    else:
        summary = (
            f"Updated 0 elements (model: {model_name}). No token stats to report."
        )

    print(summary)
    return summary


def save_elements_checkpoint(
    elements, base_filename, subdir="checkpoints", metadata=None
):
    """
    Save the current state of `elements` as JSON into a checkpoint subdirectory.
    """
    if not isinstance(elements, list):
        raise TypeError("`elements` must be a list of dictionaries.")

    os.makedirs(subdir, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"{base_filename}_{ts}.json"
    out_path = os.path.join(subdir, filename)

    payload = (
        {"metadata": metadata, "elements": elements}
        if metadata is not None
        else elements
    )

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    return out_path


def load_elements_checkpoint(
    checkpoint_path: str,
) -> Tuple[Dict[str, Optional[str]], List[Dict[str, Any]]]:
    """
    Load a checkpoint and return (metadata, elements).
    """
    with open(checkpoint_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    if isinstance(payload, list):
        return build_workflow_metadata(), payload

    if isinstance(payload, dict):
        elements = payload.get("elements")
        if not isinstance(elements, list):
            raise TypeError("Checkpoint 'elements' must be a list.")

        metadata = payload.get("metadata")
        if metadata is None:
            metadata = build_workflow_metadata()
        elif not isinstance(metadata, dict):
            raise TypeError("Checkpoint 'metadata' must be a dictionary.")
        else:
            if "docxfilename" not in metadata and "word_doc_filename" in metadata:
                metadata = dict(metadata)
                metadata["docxfilename"] = metadata.get("word_doc_filename")

            normalized_metadata = build_workflow_metadata()
            normalized_metadata.update(
                {
                    key: metadata.get(key)
                    for key in normalized_metadata
                    if key in metadata
                }
            )
            metadata = normalized_metadata

        return metadata, elements

    raise TypeError("Checkpoint JSON must be a list or a dictionary.")


def plot_token_histogram(elements: List[WorkflowRuntimeElement]) -> Tuple[pd.DataFrame, Figure]:
    """Plot histogram of element token counts and return (DataFrame, Figure)."""
    if not isinstance(elements, list):
        raise TypeError("`elements` must be a list of dictionaries.")

    df = pd.DataFrame(elements)
    if "tokens" not in df.columns:
        raise ValueError("Each element must include a 'tokens' field before plotting.")

    fig, ax = plt.subplots(figsize=(7, 4))
    ax.hist(df["tokens"], bins="auto", color="steelblue", edgecolor="black")
    ax.set_xlabel("Tokens per element")
    ax.set_ylabel("Frequency")
    ax.set_title("Distribution of Element Token Sizes (tiktoken)")
    fig.tight_layout()

    stats_msg = (
        f"Elements: {len(df)} | min={df['tokens'].min()} | "
        f"median={df['tokens'].median():.0f} | max={df['tokens'].max()} | "
        f"mean={df['tokens'].mean():.1f}"
    )
    print(stats_msg)

    return df, fig


def assign_batch_numbers(
    elements: List[WorkflowRuntimeElement],
    singleton_threshold: int,
    min_target: int,
    max_target: int,
    token_key: str = "tokens",
    batch_key: str = "batch_number",
    df: Optional[pd.DataFrame] = None,
) -> Tuple[int, List[Optional[int]]]:
    """
    Assign batch identifiers to elements based on token counts.
    Returns (number_of_batches, batch_id_list).
    """

    def _token_value(item: Mapping[str, object]) -> int:
        raw_value = item.get(token_key, 0)
        if isinstance(raw_value, bool):
            return int(raw_value)
        if isinstance(raw_value, int):
            return raw_value
        if isinstance(raw_value, float):
            return int(raw_value)
        if isinstance(raw_value, str):
            try:
                return int(raw_value)
            except ValueError:
                return 0
        return 0

    batch_id = 1
    buffer_indices: List[int] = []
    buffer_sum = 0
    batch_ids: List[Optional[int]] = [None] * len(elements)

    i = 0
    total = len(elements)
    while i < total:
        tokens = _token_value(elements[i])

        if tokens >= singleton_threshold:
            if buffer_indices:
                for idx in buffer_indices:
                    elements[idx][batch_key] = batch_id
                    batch_ids[idx] = batch_id
                batch_id += 1
                buffer_indices.clear()
                buffer_sum = 0

            elements[i][batch_key] = batch_id
            batch_ids[i] = batch_id
            batch_id += 1
            i += 1
            continue

        buffer_indices.append(i)
        buffer_sum += tokens

        if buffer_sum >= min_target:
            j = i + 1
            while j < total:
                next_tokens = _token_value(elements[j])
                if next_tokens >= singleton_threshold:
                    break
                if buffer_sum + next_tokens > max_target:
                    break
                buffer_indices.append(j)
                buffer_sum += next_tokens
                j += 1

            for idx in buffer_indices:
                elements[idx][batch_key] = batch_id
                batch_ids[idx] = batch_id
            batch_id += 1

            i = buffer_indices[-1] + 1
            buffer_indices.clear()
            buffer_sum = 0
            continue

        i += 1

    if buffer_indices:
        for idx in buffer_indices:
            elements[idx][batch_key] = batch_id
            batch_ids[idx] = batch_id
        batch_id += 1

    total_batches = batch_id - 1

    if df is not None:
        if len(df) != len(batch_ids):
            raise ValueError(
                "DataFrame length must match the number of elements when syncing batch numbers."
            )
        df[batch_key] = batch_ids

    summary = f"Assigned batch numbers to {len(elements)} elements across {total_batches} batches."
    print(summary)
    return total_batches, batch_ids

def _abbreviate_key(key: Optional[str], visible: int = 4) -> str:
    """Return a short representation of an API key for status messages."""
    if not key:
        return "missing"
    if len(key) <= visible * 2:
        return key
    return f"{key[:visible]}...{key[-visible:]}"


def _load_dotenv() -> str:
    """
    Load the .env file and report Gemini API key status.
    """
    dotenv_loaded = load_dotenv(override=True)
    if not dotenv_loaded:
        print(
            "dotenv file not found. Copy your dotenv file containing the API keys "
            "to the working directory."
        )
        return "dotenv file missing"

    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise RuntimeError("GEMINI_API_KEY not found in environment")

    status = f"Loaded API keys | Gemini={_abbreviate_key(gemini_api_key)}"
    return status


def initialize_gemini_client() -> Any:
    """
    Instantiate a Gemini client using the current Google GenAI SDK.
    """
    _load_dotenv()

    if genai is None:
        raise ImportError(
            "google.genai is not installed. Install the Google GenAI SDK to use Gemini."
        )

    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise RuntimeError("GEMINI_API_KEY not found in environment")

    # Avoid accidental fallback to a different Google API key in shared environments.
    os.environ.pop("GOOGLE_API_KEY", None)

    gemini_client = genai.Client(api_key=gemini_api_key)
    print(f"Gemini client initialized successfully: {gemini_client}")
    return gemini_client


def initialize_openai_client() -> Any:
    """
    Instantiate an OpenAI client using OPENAI_API_KEY from the current environment.
    """
    dotenv_loaded = load_dotenv(override=True)
    if not dotenv_loaded:
        print(
            "dotenv file not found. Copy your dotenv file containing the API keys "
            "to the working directory."
        )

    if OpenAI is None:
        raise ImportError(
            "openai is not installed. Install the OpenAI SDK to use evaluator helpers."
        )

    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        raise RuntimeError("OPENAI_API_KEY not found in environment")

    return OpenAI(api_key=openai_api_key)

def generate_payload(elements, batch_number):
    """
    Build a translation payload for a given batch_number from a canonical `elements` list.

    Parameters
    ----------
    elements : list[dict]
        Canonical list of element dictionaries, each containing at least:
        - "element_id"
        - "text"
        - "batch_number"
        - "element_number" (for ordering within the batch)
    batch_number : int
        The batch number to filter elements by.

    Returns
    -------
    dict
        A payload dictionary of the form:
        {
            "batch_number": <int>,
            "elements": [
                {"id": <element_id>, "text": <text>},
                ...
            ]
        }

    Raises
    ------
    ValueError
        If no elements are found for the given batch_number.
    """
    # Filter elements belonging to this batch_number
    batch_elements = [
        el for el in elements
        if el.get("batch_number") == batch_number
    ]

    if not batch_elements:
        raise ValueError(f"No elements found for batch_number={batch_number}")

    # Sort by element_number to preserve document order within the batch
    batch_elements.sort(key=lambda el: el.get("element_number", 0))

    # Build minimal payload elements (id + text only)
    payload_elements = [
        {
            "id": el["element_id"],
            "text": el["text"],
        }
        for el in batch_elements
    ]

    # Wrap in a payload object (can be extended later if needed)
    payload = {
        "batch_number": batch_number,
        "elements": payload_elements,
    }

    return payload


def _strip_json_code_fences(text: str) -> str:
    """Remove surrounding Markdown code fences from a JSON response body."""
    stripped = text.strip()
    if stripped.startswith("```"):
        lines = stripped.splitlines()

        if lines and lines[0].lstrip().startswith("```"):
            lines = lines[1:]

        if lines and lines[-1].strip().startswith("```"):
            lines = lines[:-1]

        stripped = "\n".join(lines).strip()

    return stripped


def _extract_openai_response_text(response: Any) -> str:
    """Extract plain text from an OpenAI Responses API result."""
    output_text = getattr(response, "output_text", None)
    if output_text:
        return cast(str, output_text)

    output = getattr(response, "output", None) or []
    text_parts: List[str] = []

    for item in output:
        content = getattr(item, "content", None) or []
        for part in content:
            part_text = getattr(part, "text", None)
            if part_text:
                text_parts.append(str(part_text))

    return "\n".join(text_parts).strip()


def generate_evaluator_payload(elements, batch_number):
    """
    Build an evaluator payload for one batch from the canonical `elements` list.
    """
    batch_elements = [
        el for el in elements
        if el.get("batch_number") == batch_number
        and el.get("primary_translation") is not None
        and el.get("primary_error") is None
        and el.get("final") is None
    ]

    payload_elements = [
        {
            "element_id": el["element_id"],
            "source_text": el["text"],
            "candidate_text": el["primary_translation"],
        }
        for el in batch_elements
    ]

    return {"elements": payload_elements}


def run_evaluation_pass(
    elements,
    openai_model_name,
    evaluator_system_message,
    verbose=True,
    print_status_every_n_batches=20,
):
    """
    Run the evaluator stage over all distinct batch numbers in the canonical `elements` list.
    """
    batch_numbers = sorted({el.get("batch_number") for el in elements})
    total_batches = len(batch_numbers)

    if total_batches == 0:
        if verbose:
            now = datetime.now().strftime("%H:%M:%S")
            print(f"[{now}] No batches found for evaluation.")
        return elements

    openai_client = initialize_openai_client()

    start_time = time.time()
    if verbose:
        now = datetime.now().strftime("%H:%M:%S")
        print(f"[{now}] Starting evaluation: {total_batches} batches detected.")

    for idx, batch_number in enumerate(batch_numbers, start=1):
        payload = generate_evaluator_payload(elements, batch_number)
        eligible_ids = [
            el["element_id"]
            for el in elements
            if el.get("batch_number") == batch_number
            and el.get("primary_translation") is not None
            and el.get("primary_error") is None
            and el.get("final") is None
        ]

        if not eligible_ids:
            continue

        try:
            request_json = json.dumps(payload, ensure_ascii=False)
            response = openai_client.responses.create(
                model=openai_model_name,
                input=[
                    {"role": "system", "content": evaluator_system_message},
                    {"role": "user", "content": request_json},
                ],
            )

            model_output = _extract_openai_response_text(response)
            if not model_output:
                raise ValueError("Model returned no text.")

            stripped = _strip_json_code_fences(model_output)
            if not stripped:
                raise ValueError("Empty body after stripping code fences.")

            parsed = json.loads(stripped)

        except Exception as e:
            err_msg = f"Evaluation error: {e}"
            for el in elements:
                if el.get("element_id") in eligible_ids:
                    el["evaluator_error"] = err_msg
            continue

        returned_items = parsed.get("elements", [])
        returned_by_id = {item["element_id"]: item for item in returned_items}

        for el in elements:
            if el.get("element_id") not in eligible_ids:
                continue

            returned = returned_by_id.get(el["element_id"])
            if not returned:
                el["evaluator_error"] = "Missing result for this element in model output"
                continue

            passed = returned.get("passed")
            if not isinstance(passed, bool):
                el["evaluator_error"] = "Invalid evaluator result: 'passed' must be boolean"
                continue

            feedback = returned.get("feedback")
            if passed:
                feedback = feedback or ""

            el["evaluator_ran"] = True
            el["evaluator_passed"] = passed
            el["evaluator_feedback"] = feedback
            el["evaluator_error"] = None

        if verbose and (idx % print_status_every_n_batches == 0):
            now = datetime.now().strftime("%H:%M:%S")
            print(f"[{now}] Progress: {idx}/{total_batches} batches completed.")

    if verbose:
        elapsed = time.time() - start_time
        hours = int(elapsed // 3600)
        minutes = int((elapsed % 3600) // 60)
        now = datetime.now().strftime("%H:%M:%S")
        print(
            f"[{now}] Evaluation complete: "
            f"{total_batches}/{total_batches} batches processed "
            f"in {hours:02d}:{minutes:02d}."
        )

    return elements


def run_primary_translation(
        elements, 
        gemini_client, 
        primary_system_message,
        primary_model_name, 
        verbose=True,
        print_status_every_n_batches=20,
):
    """
    Run the PRIMARY translation stage over all batches in the canonical `elements` list.
    For each batch_number:
      - Build a minimal payload using generate_payload(elements, batch_number)
      - Send the payload to the Gemini model (gemini_client)
      - Parse the model's JSON output (even if wrapped in ```json fences)
      - Write primary translation results back into the corresponding `elements`

    Parameters
    ----------
    elements : list[dict]
        Canonical list of element dictionaries, each already containing:
        - element_id
        - text
        - batch_number
        - primary_translation / primary_translation_model / primary_error fields (initially None)
    gemini_client : Any
        Gemini client instance created by initialize_gemini_client().
    primary_system_message : str
        System prompt to send with each Gemini request.
    verbose : bool, optional
        If True, print high-level progress updates (start, every 20 batches, completion).

    Returns
    -------
    elements : list[dict]
        The same canonical elements list, modified in place with:
        - primary_translation
        - primary_translation_model
        - primary_error    (if any error occurs)
    """
    # ------------------------------------------------------------------
    # 1. Discover all distinct batch numbers in order
    # ------------------------------------------------------------------
    batch_numbers = sorted({el.get("batch_number") for el in elements})
    total_batches = len(batch_numbers)

    if total_batches == 0:
        if verbose:
            now = datetime.now().strftime("%H:%M:%S")
            print(f"[{now}] No batches found for primary translation.")
        return elements

    start_time = time.time()
    if verbose:
        now = datetime.now().strftime("%H:%M:%S")
        print(f"[{now}] Starting primary translation: {total_batches} batches detected.")

    # ------------------------------------------------------------------
    # 2. Process each batch
    # ------------------------------------------------------------------
    for idx, batch_number in enumerate(batch_numbers, start=1):

        # ------------------------------------------------------------------
        # Build payload
        # ------------------------------------------------------------------
        try:
            payload = generate_payload(elements, batch_number)
        except Exception as e:
            # Record the error at the element level since this batch cannot run
            err_msg = f"Payload generation error: {e}"
            for el in elements:
                if el.get("batch_number") == batch_number:
                    el["primary_error"] = err_msg
            continue

        # ------------------------------------------------------------------
        # Call Gemini and parse JSON (with code-fence stripping)
        # ------------------------------------------------------------------
        try:
            # Serialize payload to JSON
            request_json = json.dumps(payload, ensure_ascii=False)

            request_kwargs = {
                "model": primary_model_name,
                "contents": request_json,
            }
            if genai_types is not None and hasattr(genai_types, "GenerateContentConfig"):
                request_kwargs["config"] = genai_types.GenerateContentConfig(
                    system_instruction=primary_system_message,
                )
            else:
                request_kwargs["config"] = {
                    "system_instruction": primary_system_message,
                }

            # Preserve the existing timeout target when the installed SDK supports it.
            try:
                response = gemini_client.models.generate_content(
                    **request_kwargs,
                    timeout=30,
                )
            except TypeError:
                response = gemini_client.models.generate_content(
                    **request_kwargs,
                )

            # Extract the raw text
            model_output = response.text if hasattr(response, "text") else str(response)
            if model_output is None:
                raise ValueError("Model returned no text.")

            # Strip ```json / ``` code fences if the model added them
            stripped = model_output.strip()
            if stripped.startswith("```"):
                lines = stripped.splitlines()

                # Drop opening ``` or ```json
                if lines and lines[0].lstrip().startswith("```"):
                    lines = lines[1:]

                # Drop trailing ``` if present
                if lines and lines[-1].strip().startswith("```"):
                    lines = lines[:-1]

                stripped = "\n".join(lines).strip()

            if not stripped:
                raise ValueError("Empty body after stripping code fences.")

            # Parse the cleaned JSON
            parsed = json.loads(stripped)

        except Exception as e:
            # Failure at API, timeout, network, or JSON parsing stage
            err_msg = f"Primary translation error: {e}"
            for el in elements:
                if el.get("batch_number") == batch_number:
                    el["primary_error"] = err_msg
            continue

        # ------------------------------------------------------------------
        # 3. Write primary translations back into canonical `elements`
        # ------------------------------------------------------------------
        returned_items = parsed.get("elements", [])
        returned_by_id = {item["id"]: item for item in returned_items}

        for el in elements:
            if el.get("batch_number") != batch_number:
                continue

            element_id = el["element_id"]
            returned = returned_by_id.get(element_id)

            if not returned:
                el["primary_error"] = "Missing result for this element in model output"
                continue

            el["primary_translation"] = returned.get("translated_text")
            el["primary_translation_model"] = primary_model_name
            
            el["primary_error"] = None  # Clear any previous errors

        # ------------------------------------------------------------------
        # Progress updates
        # ------------------------------------------------------------------
        if verbose and (idx % print_status_every_n_batches == 0):
            now = datetime.now().strftime("%H:%M:%S")
            print(f"[{now}] Progress: {idx}/{total_batches} batches completed.")

    # ------------------------------------------------------------------
    # 4. Completion summary
    # ------------------------------------------------------------------
    if verbose:
        elapsed = time.time() - start_time
        hours = int(elapsed // 3600)
        minutes = int((elapsed % 3600) // 60)
        now = datetime.now().strftime("%H:%M:%S")
        print(
            f"[{now}] Primary translation complete: "
            f"{total_batches}/{total_batches} batches processed "
            f"in {hours:02d}:{minutes:02d}."
        )

    return elements

def get_checkpoints_dir(base_dir: Optional[str] = None) -> str:
    """
    Return the path to the 'checkpoints' subdirectory rooted at base_dir or CWD.
    """
    root = base_dir or os.getcwd()
    checkpoints_dir = os.path.join(root, "checkpoints")
    return checkpoints_dir


def list_json_files(checkpoints_dir: Optional[str] = None) -> List[str]:
    """
    List *.json files inside the provided checkpoints_dir (defaults to 'checkpoints').
    """
    target_dir = checkpoints_dir or get_checkpoints_dir()
    if not os.path.isdir(target_dir):
        print(f"'checkpoints' directory not found at {target_dir}")
        return []

    json_files = sorted(
        fname for fname in os.listdir(target_dir) if fname.lower().endswith(".json")
    )

    if json_files:
        print("Found JSON files:")
        for fname in json_files:
            print(f"- {fname}")
    else:
        print("No JSON files found in 'checkpoints'.")

    return json_files

default_fields = {
    # Primary Translator fields
    "primary_translation": None,
    "primary_translation_model": None,
    "primary_error": None,
    # Evaluator fields
    "evaluator_ran": None,
    "evaluator_passed": None,
    "evaluator_feedback": None,
    "evaluator_error": None,
    # Fallback Translator fields
    "fallback_translation": None,
    "fallback_translation_model": None,
    "fallback_error": None,
    # Final accepted result
    "final": None,
    "final_model": None,
}


def apply_element_schema(
    elements: List[Dict[str, str]],
    schema_overrides: Optional[Dict[str, Optional[str]]] = None,
    overwrite_existing: bool = False,
) -> List[Dict[str, str]]:
    """
    Ensure each element dictionary includes the translation workflow schema keys.
    """
    if not isinstance(elements, list):
        raise TypeError("`elements` must be a list of dictionaries.")

    schema_fields = dict(default_fields)
    if schema_overrides:
        schema_fields.update(schema_overrides)

    updated = 0
    for element in elements:
        if not isinstance(element, dict):
            raise TypeError("Each element must be a dictionary.")
        for key, value in schema_fields.items():
            if overwrite_existing or key not in element:
                element[key] = value
        updated += 1

    print(
        f"Applied schema ({len(schema_fields)} fields) to {updated} elements "
        f"(overwrite_existing={overwrite_existing})."
    )
    return elements


def save_elements_checkpoint(
    elements, base_filename, subdir="checkpoints", metadata=None
):
    """
    Save the current state of `elements` as JSON into a checkpoint subdirectory.
    
    Parameters
    ----------
    elements : list[dict]
        The in-memory elements structure to serialize.
    base_filename : str
        Base name to use for the file (e.g., 'primary_translation_completed').
    subdir : str, optional
        Subdirectory (relative to current working directory) to save into.
        Default is 'checkpoints'.
    metadata : dict, optional
        Workflow metadata to save alongside the elements.
    
    Returns
    -------
    out_path : str
        Full path to the saved JSON file.
    """
    if not isinstance(elements, list):
        raise TypeError("`elements` must be a list of dictionaries.")

    # Ensure subdirectory exists
    os.makedirs(subdir, exist_ok=True)

    # Timestamp suffix: _YYYYMMDD_HHMM
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    filename = f"{base_filename}_{ts}.json"

    out_path = os.path.join(subdir, filename)

    payload = (
        {"metadata": metadata, "elements": elements}
        if metadata is not None
        else elements
    )

    # Save JSON (UTF-8, pretty-printed, no ASCII escaping)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    return out_path
